import re
import requests
from datetime import datetime
from bs4 import BeautifulSoup
from urllib.parse import urlparse, unquote
import PyPDF2
from io import BytesIO
from docx import Document
from webapp.services.log import ScrapingLog

def is_valid_url(url: str) -> bool:
    try:
        url = unquote(url.strip().rstrip('/'))
        result = urlparse(url)
        if not all([result.scheme, result.netloc]):
            if not url.startswith(('http://', 'https://')):
                url = 'https://' + url
                result = urlparse(url)
            if not all([result.scheme, result.netloc]):
                ScrapingLog.add_log(f"URL validation failed: {url} (scheme: {result.scheme}, netloc: {result.netloc})")
                return False
        return True
    except ValueError as e:
        ScrapingLog.add_log(f"URL parsing error: {url}, error: {str(e)}")
        return False

def get_format(url: str) -> str:
    parsed_url = urlparse(url)
    path = parsed_url.path.lower()
    if path.endswith('.pdf'):
        return 'PDF'
    elif path.endswith('.doc') or path.endswith('.docx'):
        return 'DOC'
    return 'HTML'

def construct_search_url(engine: str, query: str) -> str:
    query = query.replace(' ', '%20')
    if engine == "Bing":
        return f"https://www.bing.com/search?q={query}"
    elif engine == "Startpage":
        return f"https://www.startpage.com/do/dsearch?query={query}"
    elif engine == "Ecosia":
        return f"https://www.ecosia.org/search?q={query}"
    elif engine == "Yahoo":
        return f"https://search.yahoo.com/search?p={query}"
    elif engine == "DuckDuckGo":
        return f"https://duckduckgo.com/?q={query}"
    return None

def extract_description_from_response(response, format_type: str) -> str:
    try:
        if format_type == 'PDF':
            pdf_reader = PyPDF2.PdfReader(BytesIO(response.content))
            text = ''
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text[:500]
        elif format_type == 'DOC':
            doc = Document(BytesIO(response.content))
            text = ' '.join(paragraph.text for paragraph in doc.paragraphs)
            return text[:500]
        elif format_type == 'HTML':
            soup = BeautifulSoup(response.text, 'html.parser')
            paragraphs = soup.find_all('p')
            description = ' '.join(p.text for p in paragraphs)[:500]
            return description
        return response.text[:500]
    except Exception as e:
        ScrapingLog.add_log(f"Error extracting description: {str(e)}")
        return ""

def fetch_closing_keywords(db_connection):
    default_keywords = ["closing date", "deadline", "submission date", "due date"]
    try:
        cur = db_connection.cursor()
        cur.execute("SELECT keyword FROM closing_keywords;")
        keywords = [row[0] for row in cur.fetchall()]
        cur.close()
        if not keywords:
            return default_keywords
        return keywords
    except Exception as e:
        return default_keywords

def fetch_relevant_keywords(db_connection) -> list:
    default_keywords = ["tender", "rfp", "request for proposal", "procurement", "bid"]
    try:
        cur = db_connection.cursor()
        cur.execute("SELECT keyword FROM relevant_keywords;")
        keywords = [row[0] for row in cur.fetchall()]
        cur.close()
        if not keywords:
            ScrapingLog.add_log(f"No keywords found in relevant_keywords, using defaults: {default_keywords}")
            return default_keywords
        return keywords
    except Exception as e:
        ScrapingLog.add_log(f"Error fetching keywords: {str(e)}")
        return default_keywords

def clean_date_string(date_str: str) -> str:
    try:
        # Remove day names (e.g., "Monday", "Tuesday", etc.)
        date_str = re.sub(r'\b(Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday)\b', '', date_str, flags=re.IGNORECASE)
        
        # Remove time portions like "AT 10:00 AM"
        date_str = re.sub(r"\bAT\s+\d{1,2}:\d{2}\s*(AM|PM)?", "", date_str, flags=re.IGNORECASE)
        
        # Remove ordinal suffixes (e.g., "25th" -> "25")
        date_str = re.sub(r'\b(\d+)(st|nd|rd|th)\b', r'\1', date_str, flags=re.IGNORECASE)

        # Normalize multiple spaces to a single space
        date_str = re.sub(r'\s+', ' ', date_str).strip()

        # Lowercase and capitalize correctly (e.g., "march" -> "March")
        date_str = date_str.lower().title()

        return date_str
    except Exception as e:
        ScrapingLog.add_log(f"Error cleaning date string {date_str}: {str(e)}")
        return date_str

def extract_closing_dates(text: str, db_connection) -> list:
    closing_keywords = fetch_closing_keywords(db_connection)
    if not closing_keywords:
        return []

    date_formats = [
        r"(?:\w+day\s+)?\d{1,2}(?:st|nd|rd|th)?\s*(?:of\s+)?\w+\s+\d{4}",
        r"\w+day\s+\d{1,2}(?:st|nd|rd|th)?\s+\w+\s+\d{4}",
        r"\d{1,2}\s+\w+\s+\d{4}\s+AT\s+\d{1,2}:\d{2}\s+(?:AM|PM)",
        r"\d{1,2}(?:st|nd|rd|th)?\s+\w+\s+\d{4}",
        r"\d{1,2}/\d{1,2}/\d{4}",
        r"\d{4}-\d{2}-\d{2}",
        r"\w+\s+\d{1,2},\s+\d{4}",
        r"\w+\s+\d{1,2}\s+\d{4}",
        r"\d{1,2}-\w{3}-\d{4}"
    ]

    combined_pattern = "|".join(f"(?:{pattern})" for pattern in date_formats)
    pattern = re.compile(combined_pattern, re.IGNORECASE)

    closing_dates = []
    for keyword in closing_keywords:
        keyword_pattern = re.compile(rf"{re.escape(keyword)}\s*:\s*(.+?)(?:\n|$|\s{{2,}})", re.IGNORECASE)
        matches = keyword_pattern.finditer(text)
        for match in matches:
            date_str = match.group(1).strip()
            date_match = pattern.search(date_str)
            if date_match:
                closing_dates.append((date_match.group(), keyword))
                break

    if not closing_dates:
        for match in pattern.finditer(text):
            closing_dates.append((match.group(), None))

    ScrapingLog.add_log(f"Extracted dates: {closing_dates}")
    if not closing_dates:
        ScrapingLog.add_log(f"No dates found. Text sample: {text[:500]}...")

    return closing_dates

def parse_closing_date(date_str: str) -> datetime.date:
    try:
        date_str = clean_date_string(date_str)

        # Common formats to try
        formats_to_try = [
            "%d %B %Y",      # 25 March 2025
            "%dth %B %Y",    # 25 Emileth March 2025
            "%d %b %Y",      # 25 Mar 2025
            "%d/%m/%Y",      # 25/03/2025
            "%Y-%m-%d",      # 2025-03-25
            "%B %d, %Y",     # March 25, 2025
            "%d-%b-%Y",      # 25-Mar-2025
            "%B %d %Y",      # March 25 2025
            "%Y-%m-%d %H:%M:%S",  # 2025-04-29 10:00:00
            "%Y-%m-%d %H:%M"      # 2024-04-19 16:00
        ]

        # Handle special case with dot in time (e.g., '2023-09-28 11.00:00')
        date_str = date_str.replace('11.00:00', '11:00:00')

        for fmt in formats_to_try:
            try:
                return datetime.strptime(date_str, fmt).date()
            except ValueError:
                continue

        # Final fallback using dateutil's smart parser
        try:
            from dateutil import parser
            return parser.parse(date_str, fuzzy=True).date()
        except ImportError:
            ScrapingLog.add_log("dateutil.parser not available, skipping fuzzy parsing")
        except Exception as e:
            ScrapingLog.add_log(f"dateutil.parser failed to parse date: {date_str}, error: {str(e)}")

        ScrapingLog.add_log(f"Failed to parse closing date after trying all formats: {date_str}")
        return None
    except Exception as e:
        ScrapingLog.add_log(f"Error parsing date {date_str}: {str(e)}")
        return None

def is_relevant_tender(text: str, db_connection) -> tuple[bool, list]:
    keywords = fetch_relevant_keywords(db_connection)
    if not keywords:
        return False, []
    
    matched_keywords = []
    for keyword in keywords:
        if re.search(re.escape(keyword), text, re.IGNORECASE):
            matched_keywords.append(keyword)
    
    return bool(matched_keywords), matched_keywords

def insert_tender_to_db(tender_info, db_connection):
    try:
        cur = db_connection.cursor()
        
        # Check if the tender already exists based on source_url
        cur.execute("SELECT 1 FROM tenders WHERE source_url = %s;", (tender_info['source_url'],))
        exists = cur.fetchone() is not None

        # Perform the insert or update
        cur.execute("""
            INSERT INTO tenders (title, description, closing_date, source_url, status, scraped_at, format, tender_type, location)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            ON CONFLICT (source_url) DO UPDATE
            SET title = EXCLUDED.title,
                description = EXCLUDED.description,
                closing_date = EXCLUDED.closing_date,
                status = EXCLUDED.status,
                scraped_at = EXCLUDED.scraped_at,
                format = EXCLUDED.format,
                tender_type = EXCLUDED.tender_type,
                location = EXCLUDED.location;
        """, (
            tender_info['title'],
            tender_info['description'],
            tender_info['closing_date'],
            tender_info['source_url'],
            tender_info['status'],
            tender_info['scraped_at'],
            tender_info['format'],
            tender_info['tender_type'],
            tender_info['location']
        ))
        
        # Log the action
        action = "updated" if exists else "inserted"
        ScrapingLog.add_log(f"Tender {action} in database: source_url={tender_info['source_url']}")
        
        db_connection.commit()
        cur.close()
        return True
    except Exception as e:
        ScrapingLog.add_log(f"Error inserting/updating tender: {str(e)}")
        db_connection.rollback()
        return False

def extract_pdf_text(pdf_content: bytes) -> str:
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_content))
        text = ''
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        ScrapingLog.add_log(f"Error extracting PDF text: {str(e)}")
        return ""

def extract_docx_text(docx_content: bytes) -> str:
    try:
        doc = Document(BytesIO(docx_content))
        text = ' '.join(paragraph.text for paragraph in doc.paragraphs)
        return text
    except Exception as e:
        ScrapingLog.add_log(f"Error extracting DOCX text: {str(e)}")
        return ""